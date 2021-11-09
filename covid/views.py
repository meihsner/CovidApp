from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from .models import Person, City, Hospital, Laboratory
from .forms import PersonForm, CityForm, HospitalForm, LaboratoryForm, NewUserCreationForm, UpdateUserForm,\
    UpdatePasswordForm
import xlwt
from django.http import HttpResponse
from datetime import datetime
from docx import *
from django.contrib.auth.models import User
from django.contrib.auth import update_session_auth_hash


@staff_member_required
def user_registration(request):
    if request.method == 'POST':
        form_user = NewUserCreationForm(request.POST)
        if form_user.is_valid():
            form_user.save()
            return redirect(administration_panel)
    else:
        form_user = NewUserCreationForm()
    return render(request, 'registration/register.html', {'form_user': form_user, 'edit': False})


@staff_member_required
def edit_user(request, id):
    user = get_object_or_404(User, pk=id)
    form_user = UpdateUserForm(request.POST or None, instance=user)
    if form_user.is_valid():
        form_user.save()
        return redirect(administration_panel)
    return render(request, 'registration/register.html', {'form_user': form_user, 'edit': True})


@staff_member_required
def delete_user(request, id):
    user = get_object_or_404(User, pk=id)
    if request.method == "POST" and id != 1:
        user.delete()
        return redirect(administration_panel)
    return render(request, 'delete_confirm.html', {'user': user, 'delete_user': True})


@staff_member_required
def activate_deactivate_user(request, id):
    user = get_object_or_404(User, pk=id)
    if request.method == "POST" and id != 1:
        if user.is_active:
            user.is_active = False
        else:
            user.is_active = True
        user.save()
        return redirect(administration_panel)
    return render(request, 'activate_deactivate_confirm.html', {'user': user, 'is_active': user.is_active})


@staff_member_required
def administration_panel(request):
    try:
        users = User.objects.filter(username__icontains=request.GET.get('search'))
    except ValueError:
        users = User.objects.all()
    return render(request, 'administration_panel.html', {'users': users})


@login_required
def change_password(request):
    form_change_password = UpdatePasswordForm(request.user, data=request.POST)
    if form_change_password.is_valid():
        form_change_password.save()
        update_session_auth_hash(request, form_change_password.user)
        return redirect(main, 'none')
    return render(request, 'change_password.html', {'form_change_password': form_change_password})


@login_required
def main(request, sort):
    try:
        people = Person.objects.filter(surname__icontains=request.GET.get('search'))
    except ValueError:
        if sort == 'name':
            people = Person.objects.order_by('name')
        elif sort == 'surname':
            people = Person.objects.order_by('surname')
        elif sort == 'age':
            people = Person.objects.order_by('age')
        elif sort == 'city':
            people = Person.objects.order_by('city__name')
        elif sort == 'quarantine':
            people = Person.objects.order_by('-quarantine')
        elif sort == 'hospitalization':
            people = Person.objects.order_by('-hospitalization')
        elif sort == 'supervision':
            people = Person.objects.order_by('-supervision')
        elif sort == 'who_added':
            people = Person.objects.order_by('who_added')
        else:
            people = Person.objects.all()
    return render(request, 'main.html', {'people': people})


@login_required
def add_person(request):
    current_date = datetime.now().date()
    current_date = current_date.strftime('%Y-%m-%d')
    form_person = PersonForm(request.POST or None, initial={'who_added': request.user.username, 'telephone_number': '+48', 'date_of_received_information': current_date})
    if form_person.is_valid():
        form_person.save()
        return redirect(main, 'none')
    return render(request, 'add_person.html', {'form_person': form_person, 'edit': False})


@login_required
def edit_person(request, id):
    person = get_object_or_404(Person, pk=id)
    form_person = PersonForm(request.POST or None, instance=person)
    if form_person.is_valid():
        form_person.save()
        return redirect(main, 'none')
    return render(request, 'add_person.html', {'form_person': form_person, 'edit': True})


@login_required
def delete_person(request, id):
    person = get_object_or_404(Person, pk=id)
    if request.method == "POST":
        person.delete()
        return redirect(main, 'none')
    return render(request, 'delete_confirm.html', {'person': person, 'delete_user': False})


@login_required
def add_laboratory(request):
    if Laboratory.objects.count() <= 5:
        laboratories = Laboratory.objects.order_by('-id')
    else:
        laboratories = Laboratory.objects.order_by('-id')[0:6]

    form_laboratory = LaboratoryForm(request.POST or None)
    if form_laboratory.is_valid():
        form_laboratory.save()
        return redirect(main, 'none')
    return render(request, 'add_laboratory.html', {'form_laboratory': form_laboratory, 'laboratories': laboratories})


@login_required
def add_hospital(request):
    if Hospital.objects.count() <= 5:
        hospitals = Hospital.objects.order_by('-id')
    else:
        hospitals = Hospital.objects.order_by('-id')[0:6]

    form_hospital = HospitalForm(request.POST or None)
    if form_hospital.is_valid():
        form_hospital.save()
        return redirect(main, 'none')
    return render(request, 'add_hospital.html', {'form_hospital': form_hospital, 'hospitals': hospitals})


@login_required
def add_city(request):
    if City.objects.count() <= 5:
        cities = City.objects.order_by('-id')
    else:
        cities = City.objects.order_by('-id')[0:6]

    form_city = CityForm(request.POST or None)
    if form_city.is_valid():
        form_city.save()
        return redirect(main, 'none')
    return render(request, 'add_city.html', {'form_city': form_city, 'cities': cities})


@login_required
def export_data_xls(request):
    current_date = datetime.now().date()
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="PSSE {}.xls"'.format(current_date)

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Spis osób z wynikiem dodatnim')

    # Sheet header, first row
    row_num = 0
    column_width = [7, 20, 20, 7, 7, 30, 20, 30, 30, 30, 30, 30, 25, 27, 27, 25]
    for idx in range(0, len(column_width)):
        ws.col(idx).width = 256 * column_width[idx]

    font_style = xlwt.XFStyle()
    font_style.font.bold = True

    alignment = xlwt.Alignment()
    alignment.horz = xlwt.Alignment.HORZ_CENTER
    alignment.vert = xlwt.Alignment.VERT_CENTER
    alignment.wrap = 1
    borders = xlwt.Borders()
    borders.left = xlwt.Borders.THIN
    borders.right = xlwt.Borders.THIN
    borders.top = xlwt.Borders.THIN
    borders.bottom = xlwt.Borders.THIN

    font_style.alignment = alignment
    font_style.borders =borders


    columns = ['Id', 'Imię', 'Nazwisko', 'Płeć', 'Wiek', 'Miejscowość zamieszkania chorego', 'Numer telefonu', 'Data pozyskania informacji przez PSSE',
              'Data uzyskania wyniku dodatniego', 'Laboratorium wykonujące badanie', 'Miejscowość pobytu chorego', 'Źródło zakażenia',
              'Hospitalizowany', 'Nazwa szpitala', 'Objęty nadzorem', 'Poddany kwarantannie']

    for col_num in range(len(columns)):
        ws.write(row_num, col_num, columns[col_num], font_style)

    person_table = Person.objects.all().values('name', 'surname', 'gender', 'age', 'city__name', 'telephone_number', 'date_of_received_information',
              'date_of_positive_result', 'laboratory_performing_tests__name', 'whereabouts__name', 'source_of_infection',
              'hospitalization', 'hospital__name', 'supervision', 'quarantine')

    for item in person_table:
        item['date_of_positive_result'] = str(item['date_of_positive_result'])
        item['date_of_received_information'] = str(item['date_of_received_information'])

    data = [item for item in person_table]
    for item in data:
        tmp_gender = Person(gender=item['gender'])
        item['gender'] = tmp_gender.get_gender_display()

        tmp_hospitalization = Person(hospitalization=item['hospitalization'])
        item['hospitalization'] = tmp_hospitalization.get_hospitalization_display()

        tmp_supervision = Person(supervision=item['supervision'])
        item['supervision'] = tmp_supervision.get_supervision_display()

        tmp_quarantine = Person(quarantine=item['quarantine'])
        item['quarantine'] = tmp_quarantine.get_quarantine_display()

    hospitalized_font = 'font: italic  1, color red;'
    normal_font = 'font: italic  1, color black;'

    person_excel_id = 1
    for row in person_table:
        row = list(row.values())
        row.insert(0, person_excel_id)
        row_num += 1
        if row[-4] == "Tak":
            style = xlwt.easyxf(hospitalized_font)
            style.alignment = alignment
        else:
            style = xlwt.easyxf(normal_font)
            style.alignment = alignment
        for col_num in range(len(row)):
            ws.write(row_num, col_num, row[col_num], style)
        person_excel_id += 1

    wb.save(response)

    return response

@login_required
def export_notification_docx(request, id):
    person = get_object_or_404(Person, pk=id)
    current_date = datetime.now().date()

    document = Document('zawiadomienie.docx')
    paragraphs = document.paragraphs
    paragraphs[2]._p.clear()
    paragraphs[2].add_run('{} \t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\tGliwice, {}r.'.format(person.telephone_number, current_date))
    paragraphs[5]._p.clear()
    paragraphs[5].add_run('\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t{} {}'.format(person.name, person.surname))
    paragraphs[6]._p.clear()
    paragraphs[6].add_run('\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t{}'.format(person.city))
    paragraphs[7]._p.clear()
    paragraphs[7].add_run('\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t{}'.format(person.address))

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename={}_{}_zawiadomienie.docx'.format(person.name, person.surname)

    document.save(response)

    return response
